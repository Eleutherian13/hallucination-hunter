"""
File handling utilities for document parsing and processing
"""

import base64
import io
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
from dataclasses import dataclass, field

import PyPDF2
import pdfplumber


@dataclass
class DocumentMetadata:
    """Metadata for a parsed document"""
    filename: str
    file_type: str
    page_count: int
    total_characters: int
    encoding: str = "utf-8"
    additional_info: Dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Represents a parsed document with content and metadata"""
    content: str
    metadata: DocumentMetadata
    pages: List[str] = field(default_factory=list)
    

class FileHandler:
    """Handles file operations for various document types"""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx", ".html", ".md"}
    
    @classmethod
    def read_file(cls, file_path: Union[str, Path]) -> bytes:
        """Read file content as bytes"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        with open(path, "rb") as f:
            return f.read()
    
    @classmethod
    def read_text_file(cls, file_path: Union[str, Path], encoding: str = "utf-8") -> str:
        """Read text file content"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        with open(path, "r", encoding=encoding) as f:
            return f.read()
    
    @classmethod
    def save_file(cls, content: Union[str, bytes], file_path: Union[str, Path]) -> None:
        """Save content to file"""
        path = Path(file_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        
        mode = "wb" if isinstance(content, bytes) else "w"
        with open(path, mode) as f:
            f.write(content)
    
    @classmethod
    def get_file_extension(cls, file_path: Union[str, Path]) -> str:
        """Get file extension in lowercase"""
        return Path(file_path).suffix.lower()
    
    @classmethod
    def is_supported(cls, file_path: Union[str, Path]) -> bool:
        """Check if file type is supported"""
        ext = cls.get_file_extension(file_path)
        return ext in cls.SUPPORTED_EXTENSIONS
    
    @classmethod
    def decode_base64(cls, content: str) -> bytes:
        """Decode base64 encoded content"""
        return base64.b64decode(content)
    
    @classmethod
    def encode_base64(cls, content: bytes) -> str:
        """Encode content to base64"""
        return base64.b64encode(content).decode("utf-8")


class DocumentParser:
    """Parses various document formats into text content"""
    
    def __init__(self):
        self.parsers = {
            ".pdf": self._parse_pdf,
            ".txt": self._parse_text,
            ".md": self._parse_text,
            ".html": self._parse_html,
            ".docx": self._parse_docx,
        }
    
    def parse(
        self, 
        file_path: Optional[Union[str, Path]] = None,
        file_content: Optional[bytes] = None,
        file_name: Optional[str] = None,
        file_type: Optional[str] = None
    ) -> ParsedDocument:
        """
        Parse a document from file path or content bytes
        
        Args:
            file_path: Path to the file
            file_content: Raw file content as bytes
            file_name: Original filename (for type detection)
            file_type: Explicit file type (e.g., ".pdf")
        
        Returns:
            ParsedDocument with content and metadata
        """
        if file_path:
            path = Path(file_path)
            file_content = FileHandler.read_file(path)
            file_name = path.name
            file_type = path.suffix.lower()
        elif file_content:
            if not file_type and file_name:
                file_type = Path(file_name).suffix.lower()
            elif not file_type:
                raise ValueError("Must provide file_type or file_name for content parsing")
            file_name = file_name or f"document{file_type}"
        else:
            raise ValueError("Must provide either file_path or file_content")
        
        parser = self.parsers.get(file_type)
        if not parser:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        content, pages = parser(file_content)
        
        metadata = DocumentMetadata(
            filename=file_name,
            file_type=file_type,
            page_count=len(pages) if pages else 1,
            total_characters=len(content)
        )
        
        return ParsedDocument(content=content, metadata=metadata, pages=pages)
    
    def _parse_pdf(self, content: bytes) -> Tuple[str, List[str]]:
        """Parse PDF document using pdfplumber for better text extraction"""
        pages = []
        full_text = []
        
        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text() or ""
                    pages.append(text)
                    full_text.append(text)
        except Exception:
            # Fallback to PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text = page.extract_text() or ""
                pages.append(text)
                full_text.append(text)
        
        return "\n\n".join(full_text), pages
    
    def _parse_text(self, content: bytes) -> Tuple[str, List[str]]:
        """Parse plain text or markdown"""
        text = content.decode("utf-8", errors="ignore")
        return text, [text]
    
    def _parse_html(self, content: bytes) -> Tuple[str, List[str]]:
        """Parse HTML document, extracting text content"""
        from html.parser import HTMLParser
        
        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []
                self.skip_tags = {"script", "style", "head", "meta", "link"}
                self.current_skip = False
            
            def handle_starttag(self, tag, attrs):
                if tag in self.skip_tags:
                    self.current_skip = True
            
            def handle_endtag(self, tag):
                if tag in self.skip_tags:
                    self.current_skip = False
            
            def handle_data(self, data):
                if not self.current_skip:
                    text = data.strip()
                    if text:
                        self.text_parts.append(text)
        
        html_text = content.decode("utf-8", errors="ignore")
        extractor = TextExtractor()
        extractor.feed(html_text)
        
        text = " ".join(extractor.text_parts)
        return text, [text]
    
    def _parse_docx(self, content: bytes) -> Tuple[str, List[str]]:
        """Parse DOCX document"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n\n".join(paragraphs)
            return text, [text]
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")


class DocumentChunkLocator:
    """Utility to locate original positions of chunks in documents"""
    
    def __init__(self, parsed_doc: ParsedDocument):
        self.doc = parsed_doc
        self._build_index()
    
    def _build_index(self):
        """Build character offset index for pages"""
        self.page_offsets = []
        current_offset = 0
        
        for page_text in self.doc.pages:
            self.page_offsets.append(current_offset)
            current_offset += len(page_text) + 2  # +2 for "\n\n" separator
    
    def find_location(self, chunk: str) -> Optional[Dict]:
        """
        Find the location of a chunk in the original document
        
        Returns:
            Dict with page number and line range, or None if not found
        """
        content = self.doc.content
        chunk_start = content.find(chunk)
        
        if chunk_start == -1:
            return None
        
        # Find which page
        page_num = 0
        for i, offset in enumerate(self.page_offsets):
            if chunk_start >= offset:
                page_num = i + 1
            else:
                break
        
        # Find line numbers within page
        page_text = self.doc.pages[page_num - 1] if page_num <= len(self.doc.pages) else ""
        page_offset = self.page_offsets[page_num - 1] if page_num <= len(self.page_offsets) else 0
        
        local_start = chunk_start - page_offset
        lines_before = page_text[:local_start].count("\n")
        lines_in_chunk = chunk.count("\n")
        
        return {
            "page": page_num,
            "start_line": lines_before + 1,
            "end_line": lines_before + lines_in_chunk + 1,
            "char_offset": chunk_start
        }
